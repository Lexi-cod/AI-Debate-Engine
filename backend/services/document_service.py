from io import BytesIO
from pathlib import Path
from typing import TypedDict

import openpyxl
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Supported file types — now includes Excel (.xlsx / .xls)
# Per spec: PDF, DOCX, and Excel are all supported upload formats.
# ---------------------------------------------------------------------------
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls"}

MAX_FILE_BYTES   = 20 * 1024 * 1024  # 20 MB per file — prevents OOM on large uploads
MAX_CHARS_PER_FILE = 20_000           # chars extracted per file
MAX_TOTAL_CHARS    = 60_000           # total chars across all files in one request


class ParsedDocument(TypedDict):
    filename: str
    filetype: str
    content_excerpt: str


class DocumentContext(TypedDict):
    documents: list[ParsedDocument]
    combined_text: str


class DocumentServiceError(Exception):
    """Raised when uploaded files are invalid or cannot be parsed."""


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append((page.extract_text() or "").strip())
    return "\n".join(chunk for chunk in chunks if chunk).strip()


def _extract_docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    chunks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(chunks).strip()


def _extract_xlsx_text(data: bytes) -> str:
    """
    Extract text from an Excel workbook (.xlsx or .xls).

    Each sheet is labelled, then every non-empty row is tab-joined so the
    LLM receives a readable, structured representation of the spreadsheet.
    Example output:
        [Sheet: Revenue Forecast]
        Quarter    Q1      Q2      Q3      Q4
        Revenue    120000  145000  138000  162000
    """
    wb = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
    sections: list[str] = []

    for sheet in wb.worksheets:
        rows: list[str] = [f"[Sheet: {sheet.title}]"]
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append("\t".join(cells))
        if len(rows) > 1:          # skip entirely empty sheets
            sections.append("\n".join(rows))

    wb.close()
    return "\n\n".join(sections).strip()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentServiceError(
            f"Unsupported file type '{ext}' for '{filename}'. "
            f"Allowed formats: PDF, DOCX, XLSX, XLS."
        )
    return ext


# Allowed MIME types mapped to extensions — rejects renamed files
# e.g. a ZIP renamed to .pdf will have MIME application/zip and be rejected.
_ALLOWED_MIME_TYPES: dict[str, set[str]] = {
    ".pdf":  {"application/pdf"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",  # some browsers send this for docx
    },
    ".xlsx": {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/octet-stream",
    },
    ".xls":  {
        "application/vnd.ms-excel",
        "application/octet-stream",
    },
}


def _validate_mime_type(filename: str, ext: str, content_type: str | None) -> None:
    """Reject files whose MIME type does not match their extension."""
    if not content_type:
        return  # browser did not send Content-Type — skip check
    # Normalise: strip parameters like "; charset=utf-8"
    mime = content_type.split(";")[0].strip().lower()
    if not mime or mime == "application/octet-stream":
        return  # generic binary — allow, rely on parser to fail safely
    allowed = _ALLOWED_MIME_TYPES.get(ext, set())
    if allowed and mime not in allowed:
        raise DocumentServiceError(
            f"'{filename}' has MIME type '{mime}' which does not match "
            f"the expected type for {ext} files. The file may be corrupted or renamed."
        )


async def _safe_read(file: UploadFile) -> bytes:
    """Read file content with a hard size cap to prevent OOM crashes."""
    # Validate MIME type before reading the full file
    _validate_mime_type(file.filename or "", Path(file.filename or "").suffix.lower(), file.content_type)
    data = await file.read(MAX_FILE_BYTES + 1)
    if len(data) > MAX_FILE_BYTES:
        raise DocumentServiceError(
            f"'{file.filename}' exceeds the 20 MB size limit. Please upload a smaller file."
        )
    return data


def _extract_text(data: bytes, ext: str, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    try:
        if ext == ".pdf":
            return _extract_pdf_text(data)
        if ext == ".docx":
            return _extract_docx_text(data)
        if ext in (".xlsx", ".xls"):
            return _extract_xlsx_text(data)
    except DocumentServiceError:
        raise
    except Exception as exc:
        raise DocumentServiceError(f"Failed to parse '{filename}': {exc}") from exc
    return ""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def parse_uploaded_documents(files: list[UploadFile] | None) -> DocumentContext:
    if not files:
        return {"documents": [], "combined_text": ""}

    parsed_docs: list[ParsedDocument] = []
    combined_parts: list[str] = []
    running_total = 0

    for file in files:
        filename = file.filename or "unknown"
        ext = _validate_extension(filename)

        # Read with size guard
        try:
            data = await _safe_read(file)
        finally:
            await file.close()

        # Extract text
        text = _extract_text(data, ext, filename).strip()
        if not text:
            continue

        # Apply per-file and global character limits
        text = text[:MAX_CHARS_PER_FILE]
        remaining = MAX_TOTAL_CHARS - running_total
        if remaining <= 0:
            break
        text = text[:remaining]
        running_total += len(text)

        parsed_docs.append({
            "filename": filename,
            "filetype": ext.lstrip("."),
            "content_excerpt": text[:1000],
        })
        combined_parts.append(f"[{filename}]\n{text}")

    return {
        "documents": parsed_docs,
        "combined_text": "\n\n".join(combined_parts),
    }